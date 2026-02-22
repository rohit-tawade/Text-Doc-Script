try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
except Exception:
    Document = None
    Pt = None
    Inches = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    WD_TAB_ALIGNMENT = None
#!/usr/bin/env python3
"""
generator.py
Convert resume.txt into a beautiful HTML (and optional PDF) resume.
Usage:
    python converter.py resume.txt
"""

import sys, re
import os
from pathlib import Path
from xml.sax.saxutils import escape
try:
    import pythoncom  # type: ignore
    import win32com.client  # type: ignore
except Exception:
    pythoncom = None
    win32com = None
PDF_EXPORT_PATHS = {
    'deshraj sharma': Path('D:/Deshraj Sir'),
    'rohit tawade': Path('D:/Resumes'),
    'mahesh sandur': Path('D:/Mahesh Sandur'),
    'mahesh pawar': Path('D:/Mahesh Pawar'),
    'paresh kumar': Path('D:/Paresh Kumar'),
    'ashish kumar': Path('D:/Ashish Kumar'),
}

try:
    from jinja2 import Template
except Exception:
    Template = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
except Exception:
    A4 = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    TA_CENTER = None
    TA_LEFT = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    ListFlowable = None
    ListItem = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

TEMPLATE_FILE = Path(__file__).parent / 'template.jinja'


def _clean_filename_token(token):
    token = (token or '').replace('_', ' ')
    token = re.sub(r"\s+", " ", token)
    return token.strip()


def extract_filename_metadata(resume_path):
    stem = Path(resume_path).stem
    parts = [p for p in stem.split('-') if p.strip()]
    metadata = {
        'first': '',
        'last': '',
        'role': '',
        'company': ''
    }
    if not parts:
        return metadata
    metadata['first'] = _clean_filename_token(parts[0])
    if len(parts) >= 2:
        metadata['last'] = _clean_filename_token(parts[1])
    if len(parts) >= 3:
        if len(parts) >= 4:
            role_parts = [_clean_filename_token(p) for p in parts[2:-1]]
            metadata['role'] = " ".join([p for p in role_parts if p])
            metadata['company'] = _clean_filename_token(parts[-1])
        else:
            role_parts = [_clean_filename_token(p) for p in parts[2:]]
            metadata['role'] = " ".join([p for p in role_parts if p])
    return metadata


def _split_name(name):
    parts = [p for p in (name or '').replace('\n', ' ').split() if p]
    if not parts:
        return '', ''
    if len(parts) == 1:
        return parts[0], ''
    return parts[0], parts[-1]

def read_file(path):
    return Path(path).read_text(encoding='utf-8', errors='ignore').strip().splitlines()

def parse_resume(lines):
    data = {
        'name': '',
        'contact': [],
        'title': '',
        'summary': '',
        'experience': [],
        'skills': [],
        'skills_raw': '',
        'certifications': [],
        'education': [],
        'company': ''
    }
    i = 0
    # Find first non-empty line as name
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        # If the next few lines are single characters, join them
        name_line = lines[i].strip()
        # Check if the next lines are single characters (vertical name)
        name_chars = [name_line]
        j = i + 1
        while j < len(lines) and len(lines[j].strip()) == 1:
            name_chars.append(lines[j].strip())
            j += 1
        if len(name_chars) > 1:
            data['name'] = ''.join(name_chars)
            i = j
        else:
            data['name'] = name_line
            i += 1
    # Collect contact until blank
    while i < len(lines) and lines[i].strip():
        contact_line = lines[i].strip()
        lowered = contact_line.lower()
        if lowered.startswith('company name:'):
            _, _, value = contact_line.partition(':')
            if value:
                data['company'] = value.strip()
        data['contact'].append(contact_line)
        i += 1
    # Insert a single blank line after any 'Nationality:' contact entry
    if data['contact']:
        new_contacts = []
        for idx, c in enumerate(data['contact']):
            new_contacts.append(c)
            try:
                is_nationality = isinstance(c, str) and c.strip().lower().startswith('nationality:')
            except Exception:
                is_nationality = False
            if is_nationality:
                # Only add a single blank if the next item is not already blank
                next_is_blank = (idx + 1 < len(data['contact']) and data['contact'][idx+1].strip() == '')
                if not next_is_blank:
                    new_contacts.append('')
        data['contact'] = new_contacts
    # After contacts, attempt to capture a standalone title line before sections
    KNOWN_SECTION_KEYS = {
        "professional summary",
        "summary",
        "professional experience",
        "experience",
        "technical skills",
        "skills",
        "certifications",
        "certification",
        "education",
        "qualification",
    }
    # Skip blank lines between contact info and potential title
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        possible_title = lines[i].strip()
        key = possible_title.lower().rstrip(":")
        if key not in KNOWN_SECTION_KEYS and not key.startswith("certificat"):
            data['title'] = possible_title
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
    # Parse sections
    section, buffer = None, []
    def flush_buffer(sec, buf):
        text = "\n".join(buf).strip()
        if not text: 
            return
        if sec == "professional summary":
            data['summary'] = text
        elif sec == "professional experience":
            parts = [p.strip() for p in text.split("\n\n") if p.strip()]
            for p in parts:
                lines_p = [ln for ln in p.splitlines() if ln.strip()]
                if not lines_p:
                    continue
                header = lines_p[0]
                duration = None
                bullets = []
                idx = 1
                if len(lines_p) > 1 and re.search(r"\d{4}", lines_p[1]):
                    duration = lines_p[1]
                    idx = 2
                # All remaining lines: treat as bullets, regardless of starting symbol
                for ln in lines_p[idx:]:
                    bullet = ln.strip()
                    # Remove leading bullet symbols if present
                    while bullet.startswith('*') or bullet.startswith('•') or bullet.startswith('-'):
                        bullet = bullet[1:].strip()
                    if bullet:
                        bullets.append(bullet)
                # If header itself starts with a bullet symbol, treat it as a bullet, not as header
                header_clean = header.strip()
                if header_clean.startswith('*') or header_clean.startswith('•') or header_clean.startswith('-'):
                    # Move header to bullets, set header to empty
                    while header_clean.startswith('*') or header_clean.startswith('•') or header_clean.startswith('-'):
                        header_clean = header_clean[1:].strip()
                    if header_clean:
                        bullets = [header_clean] + bullets
                    header = ''
                data['experience'].append({
                    "header": header,
                    "duration": duration,
                    "bullets": bullets
                })
        elif sec == "technical skills":
            # Remove leading bullet symbols from each line
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            cleaned_lines = []
            for l in lines:
                while l.startswith('*') or l.startswith('-') or l.startswith('•'):
                    l = l[1:].strip()
                cleaned_lines.append(l)
            data['skills_raw'] = '\n'.join(cleaned_lines)
            # Split on commas, semicolons, dashes, or newlines for legacy rendering
            raw_skills = re.split(r"[-,\n;]+", data['skills_raw'])
            cleaned = []
            for s in raw_skills:
                s = s.strip()
                if not s: 
                    continue
                # Add space between joined words like PythonSQL -> Python SQL
                s = re.sub(r'([a-z])([A-Z])', r'\1 \2', s)
                cleaned.append(s)
            data['skills'] = cleaned
        elif sec and sec.startswith("certificat"):
            # Remove leading symbols from each line
            certs = []
            for l in text.splitlines():
                l = l.strip()
                while l.startswith('*') or l.startswith('-') or l.startswith('•'):
                    l = l[1:].strip()
                if l:
                    certs.append(l)
            data['certifications'] = certs
        elif sec == "education":
            # Remove leading symbols from each line
            edu = []
            for l in text.splitlines():
                l = l.strip()
                while l.startswith('*') or l.startswith('-') or l.startswith('•'):
                    l = l[1:].strip()
                if l:
                    edu.append(l)
            data['education'] = edu

    while i < len(lines):
        ln = lines[i].strip()
        key = ln.lower().rstrip(":")
        if key in ("professional summary","summary"):
            flush_buffer(section, buffer); buffer=[]; section="professional summary"
        elif key in ("professional experience","experience"):
            flush_buffer(section, buffer); buffer=[]; section="professional experience"
        elif key in ("technical skills","skills"):
            flush_buffer(section, buffer); buffer=[]; section="technical skills"
        elif key.startswith("certificat"):
            flush_buffer(section, buffer); buffer=[]; section="certifications"
        elif key in ("education","qualification"):
            flush_buffer(section, buffer); buffer=[]; section="education"
        else:
            buffer.append(ln)
        i+=1
    flush_buffer(section, buffer)
    return data

def render_html(data, out_html):
    if Template is None:
        raise RuntimeError("jinja2 is not installed; HTML rendering is unavailable.")
    if not TEMPLATE_FILE.exists():
        raise RuntimeError(f"Template file not found: {TEMPLATE_FILE}")
    tpl = Template(TEMPLATE_FILE.read_text(encoding="utf-8"))
    html = tpl.render(data=data)
    out_html.write_text(html, encoding="utf-8")
    print(f"✅ Wrote HTML: {out_html}")
# PDF generation removed per user request


def sanitize_component(text, replace_space_with_hyphen=False):
    text = (text or '').strip()
    if not text:
        return ''
    text = re.sub(r"\s+", " ", text)
    safe_chars = []
    for ch in text:
        if ch in '<>:"/\\|?*_':
            safe_chars.append(' ')
        else:
            safe_chars.append(ch)
    cleaned = ''.join(safe_chars)
    cleaned = re.sub(r"\s+", "-" if replace_space_with_hyphen else " ", cleaned).strip()
    cleaned = re.sub(r"-+", "-", cleaned)
    cleaned = cleaned.strip('-. ')
    return cleaned


def clean_bullet_text(text):
    if text is None:
        return ''
    stripped = str(text).strip()
    while stripped.startswith(('*', '-', '•')):
        stripped = stripped[1:].strip()
    return stripped


def build_contact_items(contact_lines):
    entries = []
    seen = set()
    for raw in contact_lines:
        line = str(raw).strip()
        if not line or line.lower().startswith('company name:'):
            continue
        line = clean_bullet_text(line)
        if ':' in line:
            label, value = line.split(':', 1)
            label = label.strip().title()
            value = value.strip()
            if not value:
                continue
            label_key = label.lower()
            if label_key == 'email':
                value = re.sub(r"mailto:\s*", '', value, flags=re.IGNORECASE)
                value = value.replace('<', '').replace('>', '').replace('[', '').replace(']', '')
                found_emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+", value)
                if found_emails:
                    unique_emails = []
                    for email in found_emails:
                        if email not in unique_emails:
                            unique_emails.append(email)
                    value = " / ".join(unique_emails)
            elif label_key == 'address':
                value = value.rstrip(', ')
                if value and not value.lower().endswith('india'):
                    if not value.endswith('.'):
                        value = f"{value}, India"
                    else:
                        value = f"{value} India"
            entry = {'label': label, 'label_key': label_key, 'value': value}
        else:
            entry = {'label': '', 'label_key': '', 'value': line}
        dedupe_key = (entry['label_key'], entry['value'].lower())
        if entry['value'] and dedupe_key not in seen:
            entries.append(entry)
            seen.add(dedupe_key)
    return entries


def _pdf_safe_text(text):
    text = str(text or "")
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2022": "•",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("cp1252", errors="replace").decode("cp1252")


def _pdf_escape_text(text):
    text = _pdf_safe_text(text)
    text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    return text


def _wrap_pdf_text(text, max_chars):
    text = _pdf_safe_text(text).strip()
    if not text:
        return [""]
    words = text.split()
    if not words:
        return [text]
    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if len(candidate) <= max_chars:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _render_pdf_builtin(data, pdf_path, file_metadata=None):
    """Pure-Python text PDF renderer (no reportlab/fpdf dependency)."""
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    file_metadata = file_metadata or {}

    page_w = 595.0   # A4 width in points
    page_h = 842.0   # A4 height in points
    margin_l = 40.0
    margin_r = 40.0
    margin_t = 40.0
    margin_b = 40.0
    content_w = page_w - margin_l - margin_r

    def approx_text_width(text, size, bold=False):
        # Simple approximation for built-in Helvetica metrics.
        factor = 0.56 if bold else 0.52
        return len(_pdf_safe_text(text)) * size * factor

    def line_item(text, size=10.5, bold=False, align="L", indent=0.0, gap_after=0.0):
        return {
            "type": "line",
            "text": _pdf_safe_text(text),
            "size": float(size),
            "bold": bool(bold),
            "align": align,
            "indent": float(indent),
            "gap_after": float(gap_after),
        }

    def para_item(text, size=10.5, bold=False, indent=0.0, bullet=False, gap_after=0.0):
        return {
            "type": "para",
            "text": _pdf_safe_text(text),
            "size": float(size),
            "bold": bool(bold),
            "indent": float(indent),
            "bullet": bool(bullet),
            "gap_after": float(gap_after),
        }

    def expmeta_item(left, right, size=10.5, gap_after=2.0):
        return {
            "type": "expmeta",
            "left": _pdf_safe_text(left),
            "right": _pdf_safe_text(right),
            "size": float(size),
            "gap_after": float(gap_after),
        }

    def kv_para_item(label, value, size=10.5, indent=0.0, gap_after=0.0):
        # Key/value line where the label (including colon) is bold and the value is normal.
        return {
            "type": "kvpara",
            "label": _pdf_safe_text(label),
            "value": _pdf_safe_text(value),
            "size": float(size),
            "indent": float(indent),
            "gap_after": float(gap_after),
        }

    def kv_block_item(label, value, size=10.5, indent=0.0, value_indent=14.0, gap_after=0.0):
        # Label on one line (bold), value on the next wrapped line(s).
        return {
            "type": "kvblock",
            "label": _pdf_safe_text(label),
            "value": _pdf_safe_text(value),
            "size": float(size),
            "indent": float(indent),
            "value_indent": float(value_indent),
            "gap_after": float(gap_after),
        }

    def spacer_item(height):
        return {"type": "spacer", "height": float(height)}

    def rule_item(gap_after=6.0):
        return {"type": "rule", "gap_after": float(gap_after)}

    items = []

    def add_section_heading(title):
        # Professional separation: whitespace above, bold heading, then a thin rule.
        if items:
            if items[-1]["type"] == "spacer":
                items[-1]["height"] = max(items[-1].get("height", 0.0), 18.0)
            else:
                items.append(spacer_item(18))
        items.append(line_item(title.upper(), size=11.5, bold=True, gap_after=2))
        items.append(rule_item(gap_after=6))

    candidate_name = (data.get("name") or "").strip() or "Candidate Name"
    items.append(line_item(candidate_name, size=22, bold=True, align="C", gap_after=3))

    effective_title = (data.get("title") or "").strip()
    metadata_role = (file_metadata.get("role") or "").strip()
    if metadata_role and effective_title and metadata_role.lower() == effective_title.lower():
        metadata_role = ""
    headline_bits = [bit for bit in (effective_title, metadata_role) if bit]
    if headline_bits:
        items.append(line_item(" | ".join(dict.fromkeys(headline_bits)), size=11.5, bold=True, align="C", gap_after=2))
    items.append(spacer_item(28))

    contact_items = build_contact_items(data.get("contact", []))
    if contact_items:
        label_titles = {
            "phone": "Phone",
            "email": "Email",
            "address": "Address",
            "nationality": "Nationality",
        }
        for item in contact_items:
            label_key = item.get("label_key", "")
            value = (item.get("value") or "").strip()
            if not value:
                continue
            if label_key in label_titles:
                items.append(kv_para_item(f"{label_titles[label_key]}: ", value, size=10.0, gap_after=1))
            else:
                items.append(para_item(value, size=10.0, gap_after=1))
        items.append(spacer_item(6))

    summary_text = clean_bullet_text(data.get("summary", ""))
    if summary_text:
        add_section_heading("PROFESSIONAL SUMMARY")
        items.append(para_item(summary_text, size=10.0, gap_after=2))

    experience_items = data.get("experience", [])
    if experience_items:
        add_section_heading("PROFESSIONAL EXPERIENCE")
        for role in experience_items:
            header = clean_bullet_text(role.get("header", ""))
            duration = clean_bullet_text(role.get("duration", ""))
            if header or duration:
                items.append(expmeta_item(header, duration, size=10.2, gap_after=2))

            bullets = [clean_bullet_text(item) for item in role.get("bullets", []) if clean_bullet_text(item)]
            for bullet in bullets:
                items.append(para_item(bullet, size=10.0, indent=12, bullet=True, gap_after=0.8))
            items.append(spacer_item(4))

    skills_lines = [clean_bullet_text(line) for line in data.get("skills_raw", "").splitlines() if clean_bullet_text(line)]
    if skills_lines:
        add_section_heading("SKILLS")
        for entry in skills_lines:
            if ":" in entry:
                label, value = entry.split(":", 1)
                items.append(kv_para_item(f"{label.strip()}: ", value.strip(), size=10.0, gap_after=1.2))
            else:
                items.append(para_item(entry, size=10.0, indent=12, bullet=True, gap_after=0.8))

    certifications = [clean_bullet_text(item) for item in data.get("certifications", []) if clean_bullet_text(item)]
    if certifications:
        add_section_heading("CERTIFICATIONS")
        for cert in certifications:
            items.append(para_item(cert, size=10.0, indent=12, bullet=True, gap_after=0.8))

    education_lines = [clean_bullet_text(item) for item in data.get("education", []) if clean_bullet_text(item)]
    if education_lines:
        add_section_heading("EDUCATION")
        for edu in education_lines:
            items.append(para_item(edu, size=10.0, gap_after=0.8))

    page_streams = []
    current_cmds = []
    y = page_h - margin_t

    def flush_page():
        nonlocal current_cmds, y
        if not current_cmds:
            current_cmds = []
            y = page_h - margin_t
            return
        stream = "\n".join(current_cmds).encode("cp1252", errors="replace")
        page_streams.append(stream)
        current_cmds = []
        y = page_h - margin_t

    def ensure_space(height_needed):
        nonlocal y
        if y - height_needed < margin_b:
            flush_page()

    def add_text_line(text, size, bold=False, align="L", indent=0.0):
        nonlocal y
        font_name = "/F2" if bold else "/F1"
        safe = _pdf_escape_text(text)
        line_h = max(12.0, size * 1.35)
        ensure_space(line_h)
        if align == "C":
            width = approx_text_width(text, size, bold=bold)
            x = max(margin_l, (page_w - width) / 2.0)
        elif align == "R":
            width = approx_text_width(text, size, bold=bold)
            x = max(margin_l + indent, page_w - margin_r - width)
        else:
            x = margin_l + indent
        current_cmds.append("BT")
        current_cmds.append(f"{font_name} {size:.2f} Tf")
        current_cmds.append(f"1 0 0 1 {x:.2f} {y:.2f} Tm")
        current_cmds.append(f"({safe}) Tj")
        current_cmds.append("ET")
        y -= line_h

    def add_text_segments_line(segments, size, indent=0.0):
        """Draw multiple text segments on the same line (e.g., bold label + normal value)."""
        nonlocal y
        line_h = max(12.0, size * 1.35)
        ensure_space(line_h)
        x = margin_l + indent
        for text, bold in segments:
            if not text:
                continue
            safe = _pdf_escape_text(text)
            font_name = "/F2" if bold else "/F1"
            current_cmds.append("BT")
            current_cmds.append(f"{font_name} {size:.2f} Tf")
            current_cmds.append(f"1 0 0 1 {x:.2f} {y:.2f} Tm")
            current_cmds.append(f"({safe}) Tj")
            current_cmds.append("ET")
            x += approx_text_width(text, size, bold=bold)
        y -= line_h

    def add_rule_line():
        nonlocal y
        line_h = 6.0
        ensure_space(line_h)
        rule_y = y - 1.0
        current_cmds.append("q")
        current_cmds.append("0.75 w")
        current_cmds.append(f"{margin_l:.2f} {rule_y:.2f} m")
        current_cmds.append(f"{(page_w - margin_r):.2f} {rule_y:.2f} l")
        current_cmds.append("S")
        current_cmds.append("Q")
        y -= line_h

    for item in items:
        if item["type"] == "spacer":
            ensure_space(item["height"])
            y -= item["height"]
            continue

        if item["type"] == "line":
            add_text_line(
                item["text"],
                item["size"],
                bold=item["bold"],
                align=item["align"],
                indent=item["indent"],
            )
            if item.get("gap_after", 0):
                ensure_space(item["gap_after"])
                y -= item["gap_after"]
            continue

        if item["type"] == "rule":
            add_rule_line()
            if item.get("gap_after", 0):
                ensure_space(item["gap_after"])
                y -= item["gap_after"]
            continue

        if item["type"] == "expmeta":
            size = item["size"]
            left = item.get("left", "")
            right = item.get("right", "")
            line_h = max(12.0, size * 1.35)
            ensure_space(line_h)
            if left and right:
                left_w = approx_text_width(left, size, bold=True)
                right_w = approx_text_width(right, size, bold=False)
                if left_w + right_w + 18 <= content_w:
                    add_text_line(left, size=size, bold=True, align="L")
                    y += line_h  # draw right text on same visual line
                    add_text_line(right, size=size, bold=False, align="R")
                else:
                    add_text_line(left, size=size, bold=True, align="L")
                    if right:
                        add_text_line(right, size=max(9.5, size - 0.2), bold=False, align="L", indent=12)
            elif left:
                add_text_line(left, size=size, bold=True, align="L")
            elif right:
                add_text_line(right, size=size, bold=False, align="R")
            if item.get("gap_after", 0):
                ensure_space(item["gap_after"])
                y -= item["gap_after"]
            continue

        if item["type"] == "kvpara":
            size = item["size"]
            indent = item.get("indent", 0.0)
            label = item.get("label", "")
            value = item.get("value", "")
            label_width = approx_text_width(label, size, bold=True)
            value_indent = indent + label_width
            max_chars = max(20, int((content_w - value_indent) / (size * 0.52)))
            wrapped = _wrap_pdf_text(value, max_chars) if value else []
            if wrapped:
                add_text_segments_line([(label, True), (wrapped[0], False)], size=size, indent=indent)
                for segment in wrapped[1:]:
                    add_text_line(segment, size=size, bold=False, align="L", indent=value_indent)
            else:
                add_text_segments_line([(label, True)], size=size, indent=indent)
            if item.get("gap_after", 0):
                ensure_space(item["gap_after"])
                y -= item["gap_after"]
            continue

        if item["type"] == "kvblock":
            size = item["size"]
            indent = item.get("indent", 0.0)
            value_indent = item.get("value_indent", 14.0)
            label = item.get("label", "")
            value = item.get("value", "")
            add_text_line(label, size=size, bold=True, align="L", indent=indent)
            if value:
                max_chars = max(20, int((content_w - value_indent) / (size * 0.52)))
                wrapped = _wrap_pdf_text(value, max_chars)
                for segment in wrapped:
                    add_text_line(segment, size=size, bold=False, align="L", indent=value_indent)
            if item.get("gap_after", 0):
                ensure_space(item["gap_after"])
                y -= item["gap_after"]
            continue

        # Paragraph (wrapped)
        size = item["size"]
        bold = item["bold"]
        indent = item.get("indent", 0.0)
        bullet = item.get("bullet", False)
        line_h = max(12.0, size * 1.35)
        max_chars = max(20, int((content_w - indent - (10 if bullet else 0)) / (size * 0.52)))
        wrapped = _wrap_pdf_text(item["text"], max_chars)
        if bullet:
            bullet_prefix = "• "
            bullet_prefix_w = approx_text_width(bullet_prefix, size, bold=False)
            if wrapped:
                add_text_segments_line([(bullet_prefix, False), (wrapped[0], bold)], size=size, indent=indent)
                for segment in wrapped[1:]:
                    add_text_line(segment, size=size, bold=bold, align="L", indent=indent + bullet_prefix_w)
            else:
                add_text_line(bullet_prefix, size=size, bold=False, align="L", indent=indent)
        else:
            for segment in wrapped:
                add_text_line(segment, size=size, bold=bold, align="L", indent=indent)
        if item.get("gap_after", 0):
            ensure_space(item["gap_after"])
            y -= item["gap_after"]

    flush_page()
    if not page_streams:
        page_streams = [b""]

    objects = [b""]

    def alloc_obj():
        objects.append(b"")
        return len(objects) - 1

    def set_obj(obj_id, data):
        if isinstance(data, str):
            data = data.encode("cp1252", errors="replace")
        objects[obj_id] = data

    catalog_id = alloc_obj()
    pages_id = alloc_obj()
    font_regular_id = alloc_obj()
    font_bold_id = alloc_obj()

    page_ids = []
    content_ids = []
    for _ in page_streams:
        page_ids.append(alloc_obj())
        content_ids.append(alloc_obj())

    set_obj(font_regular_id, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>")
    set_obj(font_bold_id, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold /Encoding /WinAnsiEncoding >>")

    for idx, stream_bytes in enumerate(page_streams):
        page_id = page_ids[idx]
        content_id = content_ids[idx]
        page_obj = (
            f"<< /Type /Page /Parent {pages_id} 0 R "
            f"/MediaBox [0 0 {page_w:.0f} {page_h:.0f}] "
            f"/Resources << /Font << /F1 {font_regular_id} 0 R /F2 {font_bold_id} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        )
        set_obj(page_id, page_obj)
        content_obj = (
            f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("latin-1")
            + stream_bytes
            + b"\nendstream"
        )
        set_obj(content_id, content_obj)

    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    set_obj(pages_id, f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>")
    set_obj(catalog_id, f"<< /Type /Catalog /Pages {pages_id} 0 R >>")

    out = bytearray()
    out.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0] * len(objects)

    for obj_id in range(1, len(objects)):
        offsets[obj_id] = len(out)
        out.extend(f"{obj_id} 0 obj\n".encode("latin-1"))
        out.extend(objects[obj_id])
        out.extend(b"\nendobj\n")

    xref_offset = len(out)
    out.extend(f"xref\n0 {len(objects)}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for obj_id in range(1, len(objects)):
        out.extend(f"{offsets[obj_id]:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        (
            f"trailer\n<< /Size {len(objects)} /Root {catalog_id} 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("latin-1")
    )

    pdf_path.write_bytes(out)
    return pdf_path


def _render_pdf_fpdf(data, pdf_path, file_metadata=None):
    if FPDF is None:
        raise RuntimeError("fpdf2 is not installed; FPDF PDF generation is unavailable.")

    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    file_metadata = file_metadata or {}

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(12, 12, 12)
    pdf.add_page()

    def write_line(text, size=10.5, style="", align="L", ln=True):
        pdf.set_font("Helvetica", style=style, size=size)
        pdf.cell(0, 6, _pdf_safe_text(text), ln=1 if ln else 0, align=align)

    def write_para(text, size=10.5, style=""):
        pdf.set_font("Helvetica", style=style, size=size)
        pdf.multi_cell(0, 6, _pdf_safe_text(text))

    candidate_name = (data.get("name") or "").strip() or "Candidate Name"
    pdf.set_font("Helvetica", style="B", size=20)
    pdf.cell(0, 10, _pdf_safe_text(candidate_name), ln=1, align="C")

    effective_title = (data.get("title") or "").strip()
    metadata_role = (file_metadata.get("role") or "").strip()
    if metadata_role and effective_title and metadata_role.lower() == effective_title.lower():
        metadata_role = ""
    headline_bits = [bit for bit in (effective_title, metadata_role) if bit]
    if headline_bits:
        pdf.set_font("Helvetica", style="B", size=11)
        pdf.cell(0, 7, _pdf_safe_text(" | ".join(dict.fromkeys(headline_bits))), ln=1, align="C")
    pdf.ln(1)

    contact_items = build_contact_items(data.get("contact", []))
    if contact_items:
        label_titles = {
            "phone": "Phone",
            "email": "Email",
            "address": "Address",
            "nationality": "Nationality",
        }
        for item in contact_items:
            label_key = item.get("label_key", "")
            value = (item.get("value") or "").strip()
            if not value:
                continue
            if label_key in label_titles:
                write_para(f"{label_titles[label_key]}: {value}", size=10.5, style="")
            else:
                write_para(value, size=10.5, style="")
        pdf.ln(2)

    summary_text = clean_bullet_text(data.get("summary", ""))
    if summary_text:
        write_line("PROFESSIONAL SUMMARY", size=11, style="B")
        write_para(summary_text)
        pdf.ln(1)

    experience_items = data.get("experience", [])
    if experience_items:
        write_line("PROFESSIONAL EXPERIENCE", size=11, style="B")
        for role in experience_items:
            header = clean_bullet_text(role.get("header", ""))
            duration = clean_bullet_text(role.get("duration", ""))
            if header and duration:
                write_para(f"{header} | {duration}", size=10.5, style="B")
            elif header:
                write_para(header, size=10.5, style="B")
            elif duration:
                write_para(duration, size=10.5, style="")

            bullets = [clean_bullet_text(item) for item in role.get("bullets", []) if clean_bullet_text(item)]
            for bullet in bullets:
                write_para(f"- {bullet}")
            pdf.ln(1)

    skills_lines = [clean_bullet_text(line) for line in data.get("skills_raw", "").splitlines() if clean_bullet_text(line)]
    if skills_lines:
        write_line("SKILLS", size=11, style="B")
        for entry in skills_lines:
            if ":" in entry:
                write_para(entry)
            else:
                write_para(f"- {entry}")
        pdf.ln(1)

    certifications = [clean_bullet_text(item) for item in data.get("certifications", []) if clean_bullet_text(item)]
    if certifications:
        write_line("CERTIFICATIONS", size=11, style="B")
        for cert in certifications:
            write_para(f"- {cert}")
        pdf.ln(1)

    education_lines = [clean_bullet_text(item) for item in data.get("education", []) if clean_bullet_text(item)]
    if education_lines:
        write_line("EDUCATION", size=11, style="B")
        for edu in education_lines:
            write_para(edu)

    pdf.output(str(pdf_path))
    return pdf_path


def _derive_output_stem(data, file_metadata, fallback_stem):
    # Mandatory filename format: name-role-companyname.pdf
    # Prefer parsed resume content, then fall back to source filename metadata.
    name_value = (data.get('name') or '').strip()
    if not name_value:
        first = (file_metadata.get('first') or '').strip()
        last = (file_metadata.get('last') or '').strip()
        name_value = " ".join(part for part in (first, last) if part)

    role_value = (data.get('title') or '').strip() or (file_metadata.get('role') or '').strip()
    company_value = (data.get('company') or '').strip() or (file_metadata.get('company') or '').strip()

    components = []
    for value in (name_value, role_value, company_value):
        safe = sanitize_component(value, replace_space_with_hyphen=True)
        if safe:
            components.append(safe)

    if components:
        return "-".join(components)

    return sanitize_component(fallback_stem, replace_space_with_hyphen=True) or fallback_stem


def convert_text_to_resume(resume_text, output_dir=None, source_hint=None):
    """Convert resume text into DOCX/PDF outputs and return resulting file path."""
    if isinstance(resume_text, (list, tuple)):
        lines = [str(line) for line in resume_text]
    else:
        lines = str(resume_text).splitlines()
    data = parse_resume(lines)
    resolved_hint = None
    if source_hint:
        try:
            resolved_hint = Path(source_hint)
        except Exception:
            resolved_hint = None
    file_metadata = extract_filename_metadata(resolved_hint or (source_hint or '')) if (source_hint or resolved_hint) else {
        'first': '',
        'last': '',
        'role': '',
        'company': ''
    }
    fallback_stem = (resolved_hint.stem if resolved_hint else str(source_hint or 'resume')) or 'resume'
    if output_dir is None:
        if resolved_hint and resolved_hint.parent.exists():
            output_dir = resolved_hint.parent
        else:
            output_dir = Path.cwd()
    else:
        output_dir = Path(output_dir)
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        raise RuntimeError(f"Cannot create output directory {output_dir}: {exc}")
    filename_stem = _derive_output_stem(data, file_metadata, fallback_stem)
    docx_path = output_dir / f"{filename_stem}.docx"
    result_path = render_docx(data, docx_path, file_metadata)
    return result_path


def render_pdf(data, pdf_path, file_metadata=None):
    try:
        return _render_pdf_builtin(data, pdf_path, file_metadata)
    except Exception as exc:
        print(f"⚠ Built-in PDF renderer failed, trying library fallback: {exc}")
    if FPDF is not None:
        return _render_pdf_fpdf(data, pdf_path, file_metadata)
    if SimpleDocTemplate is None or ParagraphStyle is None:
        raise RuntimeError("reportlab is not installed; PDF generation is unavailable.")

    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    file_metadata = file_metadata or {}
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        name='ResumeName',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    headline_style = ParagraphStyle(
        name='ResumeHeadline',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        name='ResumeSection',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        name='ResumeBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=3,
    )
    meta_style = ParagraphStyle(
        name='ResumeMeta',
        parent=body_style,
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        spaceAfter=2,
    )

    story = []
    candidate_name = (data.get('name') or '').strip() or 'Candidate Name'
    story.append(Paragraph(escape(candidate_name), name_style))

    effective_title = (data.get('title') or '').strip()
    metadata_role = (file_metadata.get('role') or '').strip()
    if metadata_role and effective_title and metadata_role.lower() == effective_title.lower():
        metadata_role = ''
    headline_bits = [bit for bit in (effective_title, metadata_role) if bit]
    if headline_bits:
        tagline = "  |  ".join(dict.fromkeys(headline_bits))
        story.append(Paragraph(escape(tagline), headline_style))

    contact_items = build_contact_items(data.get('contact', []))
    if contact_items:
        label_titles = {
            'phone': 'Phone',
            'email': 'Email',
            'address': 'Address',
            'nationality': 'Nationality',
        }
        for item in contact_items:
            label_key = item.get('label_key', '')
            display_value = escape(item.get('value', ''))
            if not display_value:
                continue
            if label_key in label_titles:
                label = escape(label_titles[label_key])
                story.append(Paragraph(f"<b>{label}:</b> {display_value}", body_style))
            else:
                story.append(Paragraph(display_value, body_style))
        story.append(Spacer(1, 6))

    summary_text = clean_bullet_text(data.get('summary', ''))
    if summary_text:
        story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
        story.append(Paragraph(escape(summary_text), body_style))

    experience_items = data.get('experience', [])
    if experience_items:
        story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
        for role in experience_items:
            header = clean_bullet_text(role.get('header', ''))
            duration = clean_bullet_text(role.get('duration', ''))
            if header or duration:
                header_bits = []
                if header:
                    header_bits.append(f"<b>{escape(header)}</b>")
                if duration:
                    header_bits.append(f"<font color='#444444'>{escape(duration)}</font>")
                story.append(Paragraph(" &nbsp; | &nbsp; ".join(header_bits), meta_style))
            bullets = [clean_bullet_text(item) for item in role.get('bullets', []) if clean_bullet_text(item)]
            if bullets:
                bullet_items = [ListItem(Paragraph(escape(item), body_style), leftIndent=10) for item in bullets]
                story.append(
                    ListFlowable(
                        bullet_items,
                        bulletType='bullet',
                        start='circle',
                        leftIndent=10,
                        bulletFontName='Helvetica',
                        bulletFontSize=8,
                    )
                )
                story.append(Spacer(1, 2))

    skills_lines = [clean_bullet_text(line) for line in data.get('skills_raw', '').splitlines() if clean_bullet_text(line)]
    if skills_lines:
        story.append(Paragraph('SKILLS', heading_style))
        for entry in skills_lines:
            if ':' in entry:
                label, value = entry.split(':', 1)
                story.append(Paragraph(f"<b>{escape(label.strip())}:</b> {escape(value.strip())}", body_style))
            else:
                story.append(Paragraph(f"• {escape(entry)}", body_style))

    certifications = [clean_bullet_text(item) for item in data.get('certifications', []) if clean_bullet_text(item)]
    if certifications:
        story.append(Paragraph('CERTIFICATIONS', heading_style))
        cert_items = [ListItem(Paragraph(escape(cert), body_style), leftIndent=10) for cert in certifications]
        story.append(ListFlowable(cert_items, bulletType='bullet', leftIndent=10, bulletFontSize=8))

    education_lines = [clean_bullet_text(item) for item in data.get('education', []) if clean_bullet_text(item)]
    if education_lines:
        story.append(Paragraph('EDUCATION', heading_style))
        for edu in education_lines:
            story.append(Paragraph(escape(edu), body_style))

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    doc.build(story)
    return pdf_path


def convert_text_to_pdf(resume_text, output_pdf_path=None, source_hint=None):
    """Convert resume text into a PDF output and return resulting PDF path."""
    if isinstance(resume_text, (list, tuple)):
        lines = [str(line) for line in resume_text]
    else:
        lines = str(resume_text).splitlines()
    data = parse_resume(lines)

    resolved_hint = None
    if source_hint:
        try:
            resolved_hint = Path(source_hint)
        except Exception:
            resolved_hint = None
    file_metadata = extract_filename_metadata(resolved_hint or (source_hint or '')) if (source_hint or resolved_hint) else {
        'first': '',
        'last': '',
        'role': '',
        'company': ''
    }
    fallback_stem = (resolved_hint.stem if resolved_hint else str(source_hint or 'resume')) or 'resume'
    filename_stem = _derive_output_stem(data, file_metadata, fallback_stem)

    if output_pdf_path is None:
        if resolved_hint and resolved_hint.parent.exists():
            output_dir = resolved_hint.parent
        else:
            output_dir = Path.cwd()
        pdf_path = Path(output_dir) / f"{filename_stem}.pdf"
    else:
        output_pdf_path = Path(output_pdf_path)
        if output_pdf_path.exists() and output_pdf_path.is_dir():
            pdf_path = output_pdf_path / f"{filename_stem}.pdf"
        elif output_pdf_path.suffix.lower() == '.pdf':
            pdf_path = output_pdf_path.parent / f"{filename_stem}.pdf"
        elif output_pdf_path.suffix:
            pdf_path = output_pdf_path.parent / f"{filename_stem}.pdf"
        else:
            pdf_path = output_pdf_path / f"{filename_stem}.pdf"

    try:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        raise RuntimeError(f"Cannot create output directory {pdf_path.parent}: {exc}")

    return render_pdf(data, pdf_path, file_metadata)


def try_export_docx_to_pdf(docx_path, candidate_name=''):
    if win32com is None:
        return False, "pywin32 not installed"
    docx_path = Path(docx_path).resolve()
    if not docx_path.exists():
        return False, f"DOCX missing at {docx_path}"
    target_dir = PDF_EXPORT_PATHS.get(candidate_name.strip().lower())
    if target_dir:
        target_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = target_dir / docx_path.with_suffix('.pdf').name
    else:
        pdf_path = docx_path.with_suffix('.pdf')
    word = None
    try:
        if pythoncom is not None:
            pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path), ReadOnly=True)
        try:
            doc.ExportAsFixedFormat(str(pdf_path), 17)
        finally:
            doc.Close(False)
        return True, pdf_path
    except Exception as exc:
        return False, exc
    finally:
        if word is not None:
            word.Quit()
        if pythoncom is not None:
            pythoncom.CoUninitialize()


def render_docx(data, docx_path, file_metadata=None):
    import os
    if Document is None:
        raise RuntimeError("python-docx is not installed; DOCX generation is unavailable.")

    # Remove stale DOCX so Word exports never collide with an open handle.
    try:
        if os.path.exists(docx_path):
            os.remove(docx_path)
    except PermissionError:
        print(f"❌ Permission denied: {docx_path}. Please close the file and try again.")
        return None

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    base_style = doc.styles['Normal']
    base_style.font.name = 'Calibri'
    base_style.font.size = Pt(12)
    base_style.paragraph_format.line_spacing = 1.15
    base_style.paragraph_format.space_after = Pt(4)

    try:
        bullet_style = doc.styles['List Bullet']
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(12)
    except KeyError:
        bullet_style = None

    accent = RGBColor(0, 0, 0)
    muted = RGBColor(96, 96, 96)

    def add_section_heading(title: str):
        heading_text = title.strip()
        if not heading_text:
            return None
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(heading_text.upper())
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = accent
        return p

    def clean_bullet_text(text) -> str:
        if text is None:
            return ''
        stripped = str(text).strip()
        while stripped.startswith(('*', '-', '•')):
            stripped = stripped[1:].strip()
        return stripped

    def build_contact_items(contact_lines: list[str]) -> list[dict[str, str]]:
        entries: list[dict[str, str]] = []
        seen: set[tuple[str, str]] = set()
        for raw in contact_lines:
            line = raw.strip()
            if not line or line.lower().startswith('company name:'):
                continue
            line = clean_bullet_text(line)
            if ':' in line:
                label, value = line.split(':', 1)
                label = label.strip().title()
                value = value.strip()
                if not value:
                    continue
                label_key = label.lower()
                if label_key == 'email':
                    value = re.sub(r"mailto:\s*", '', value, flags=re.IGNORECASE)
                    value = value.replace('<', '').replace('>', '').replace('[', '').replace(']', '')
                    found_emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+", value)
                    if found_emails:
                        unique_emails: list[str] = []
                        for email in found_emails:
                            if email not in unique_emails:
                                unique_emails.append(email)
                        value = " / ".join(unique_emails)
                elif label_key == 'address':
                    value = value.rstrip(', ')
                    if not value.lower().endswith('india'):
                        if value and not value.endswith('.'):
                            value = f"{value}, India"
                        else:
                            value = f"{value} India"
                entry = {'label': label, 'label_key': label_key, 'value': value}
            else:
                entry = {'label': '', 'label_key': '', 'value': line}
            dedupe_key = (entry['label_key'], entry['value'].lower())
            if entry['value'] and dedupe_key not in seen:
                entries.append(entry)
                seen.add(dedupe_key)
        return entries

    candidate_name = (data.get('name') or '').strip() or 'Candidate Name'
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(6)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(candidate_name)
    name_run.font.bold = True
    name_run.font.size = Pt(28)
    name_run.font.color.rgb = accent

    effective_title = (data.get('title') or '').strip()
    file_metadata = file_metadata or {}
    metadata_role = (file_metadata.get('role') or '').strip()
    if metadata_role and effective_title and metadata_role.lower() == effective_title.lower():
        metadata_role = ''
    headline_bits = [bit for bit in (effective_title, metadata_role) if bit]
    if headline_bits:
        tagline = "  |  ".join(dict.fromkeys(headline_bits))
        tagline_para = doc.add_paragraph(tagline)
        tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_para.paragraph_format.space_after = Pt(6)
        tagline_run = tagline_para.runs[0]
        tagline_run.font.color.rgb = accent
        tagline_run.font.size = Pt(12)
        tagline_run.font.bold = True

    contact_items = build_contact_items(data.get('contact', []))
    if contact_items:
        label_titles = {
            'phone': 'Phone',
            'email': 'Email',
            'address': 'Address',
            'nationality': 'Nationality',
        }
        for item in contact_items:
            label_key = item.get('label_key', '')
            display_value = item.get('value', '')
            if not display_value:
                continue
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(3)
            if label_key in label_titles:
                label_run = paragraph.add_run(f"{label_titles[label_key]}: ")
                label_run.font.bold = True
                label_run.font.color.rgb = accent
                label_run.font.size = Pt(12)
                value_run = paragraph.add_run(display_value)
                value_run.font.size = Pt(12)
                value_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                text_run = paragraph.add_run(display_value)
                text_run.font.size = Pt(12)
                text_run.font.color.rgb = RGBColor(0, 0, 0)
        spacer_para = doc.add_paragraph()
        spacer_para.paragraph_format.space_after = Pt(12)

    summary_text = clean_bullet_text(data.get('summary', ''))
    if summary_text:
        add_section_heading('Professional Summary')
        summary_para = doc.add_paragraph(summary_text)
        summary_para.paragraph_format.space_after = Pt(8)

    experience_items = data.get('experience', [])
    if experience_items:
        add_section_heading('Professional Experience')
        for role in experience_items:
            header = clean_bullet_text(role.get('header', ''))
            duration = clean_bullet_text(role.get('duration', ''))
            if header or duration:
                header_para = doc.add_paragraph()
                header_para.paragraph_format.space_before = Pt(8)
                header_para.paragraph_format.space_after = Pt(2)
                tab_stops = header_para.paragraph_format.tab_stops
                tab_stops.clear_all()
                tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
                if header:
                    header_run = header_para.add_run(header)
                    header_run.font.bold = True
                if duration:
                    header_para.add_run('\t')
                    duration_run = header_para.add_run(duration)
                    duration_run.font.color.rgb = RGBColor(0, 0, 0)
            bullets = role.get('bullets', [])
            for bullet in bullets:
                bullet_text = clean_bullet_text(bullet)
                if not bullet_text:
                    continue
                bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
                bullet_para.paragraph_format.space_after = Pt(1)

    skills_lines = [clean_bullet_text(line) for line in data.get('skills_raw', '').splitlines() if clean_bullet_text(line)]
    if skills_lines:
        add_section_heading('Skills')
        for entry in skills_lines:
            if ':' in entry:
                label, value = entry.split(':', 1)
                skill_para = doc.add_paragraph()
                skill_para.paragraph_format.space_after = Pt(1)
                label_run = skill_para.add_run(f"{label.strip()}: ")
                label_run.font.bold = True
                skill_para.add_run(value.strip())
            else:
                doc.add_paragraph(entry, style='List Bullet')

    certifications = [clean_bullet_text(item) for item in data.get('certifications', []) if clean_bullet_text(item)]
    if certifications:
        add_section_heading('Certifications')
        for cert in certifications:
            cert_para = doc.add_paragraph(cert, style='List Bullet')
            cert_para.paragraph_format.space_after = Pt(1)

    education_lines = [clean_bullet_text(item) for item in data.get('education', []) if clean_bullet_text(item)]
    if education_lines:
        add_section_heading('Education')
        for edu in education_lines:
            edu_para = doc.add_paragraph(edu)
            edu_para.paragraph_format.space_after = Pt(2)

    doc.save(docx_path)

    pdf_result = try_export_docx_to_pdf(docx_path, data.get('name', ''))
    generated_pdf = None
    if pdf_result[0]:
        generated_pdf = pdf_result[1]
        print(f"✅ Created PDF via Word: {generated_pdf}")
    else:
        reason = pdf_result[1]
        if isinstance(reason, str) and "pywin32" in reason.lower():
            print("ℹ Skipping PDF export (pywin32 not installed). Install pywin32 for Word-based PDF generation.")
        elif reason not in (None, ''):
            print(f"⚠ Could not create PDF via Word: {reason}")

    if not generated_pdf:
        fallback_pdf_path = Path(docx_path).with_suffix('.pdf')
        try:
            generated_pdf = render_pdf(data, fallback_pdf_path, file_metadata)
            print(f"✅ Created PDF via ReportLab: {generated_pdf}")
        except Exception as exc:
            print(f"⚠ Could not create PDF via ReportLab: {exc}")

    if generated_pdf:
        try:
            os.remove(docx_path)
        except Exception as exc:
            print(f"⚠ Could not remove intermediate DOCX: {exc}")
        if hasattr(os, "startfile"):
            try:
                os.startfile(str(generated_pdf))
            except Exception as exc:
                print(f"⚠ Could not open PDF automatically: {exc}")
        return Path(generated_pdf)

    print(f"✅ Created DOCX: {docx_path}")
    return Path(docx_path)

def main():
    if len(sys.argv)<2:
        print("Usage: python converter.py resume.txt")
        sys.exit(1)
    resume_path = Path(sys.argv[1])
    if not resume_path.exists():
        print("❌ File not found:", resume_path); sys.exit(1)
    resume_text = resume_path.read_text(encoding='utf-8', errors='ignore')
    result_path = convert_text_to_resume(resume_text, resume_path.parent, resume_path)
    if result_path:
        print(f"✅ Output: {result_path}")

if __name__ == "__main__":
    main()
